import pandas as pd
import os
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet

class ReportExporter:

    def export_excel(self, results):

        df = pd.DataFrame(results)

        filepath = "data/outputs/match_scores.xlsx"

        df.to_excel(filepath, index=False)

        print(f"Excel exported: {filepath}")



    def export_word(self, report_text, filename):

        doc = Document()

        doc.add_heading("Candidate Hiring Report", level=1)

        for line in report_text.split("\n"):
            doc.add_paragraph(line)

        filepath = f"data/outputs/{filename}.docx"

        doc.save(filepath)

        print(f"Word report exported: {filepath}")



    def export_pdf(self, report_text, filename):

        filepath = f"data/outputs/{filename}.pdf"

        doc = SimpleDocTemplate(filepath)

        styles = getSampleStyleSheet()

        content = []

        for line in report_text.split("\n"):
            content.append(Paragraph(line, styles["BodyText"]))

        doc.build(content)

        print(f"PDF exported: {filepath}")

    def export_html(self, results):

        df = pd.DataFrame(results)

        filepath = "data/outputs/dashboard.html"

        df.to_html(filepath, index=False)

        print(f"HTML dashboard exported: {filepath}")

